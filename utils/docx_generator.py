from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def clean_text(text: str):
    if not text:
        return ""

    # 彻底清理 markdown 符号
    text = text.replace("*", "")
    text = text.replace("-", "")
    text = text.replace("#", "")
    text = text.replace("`", "")
    return text.strip()


def set_font(run, bold=False, size=11):
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(size)
    run.bold = bold


def generate_docx(doc_data, product_name):
    doc = Document()

    # ===== 全局字体 =====
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(11)

    # ===== 封面 =====
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"{product_name} 产品需求文档（PRD）")
    set_font(run, bold=True, size=20)

    doc.add_paragraph("\n")

    # ===== 正文 =====
    for section in doc_data:

        # 一级标题
        h1 = doc.add_heading(section["title"], level=1)

        content = section.get("content", "")

        if not content:
            doc.add_paragraph("暂无内容")
            continue

        lines = content.split("\n")

        for line in lines:
            line = clean_text(line.strip())

            if not line:
                continue

            # 二级标题（支持 ###）
            if line.startswith("###"):
                doc.add_heading(line.replace("###", "").strip(), level=2)
                continue

            # 列表项（支持 - / *）
            if line.startswith("-") or line.startswith("*"):
                p = doc.add_paragraph(style='List Bullet')
                run = p.add_run(line.lstrip("-* ").strip())
                set_font(run)

                p.paragraph_format.left_indent = Inches(0.3)
                continue

            # 普通段落
            p = doc.add_paragraph()
            run = p.add_run(line)
            set_font(run)

            # 首行缩进 + 段间距
            p.paragraph_format.first_line_indent = Inches(0.3)
            p.paragraph_format.space_after = Pt(6)

    file_path = f"{product_name}_PRD.docx"
    doc.save(file_path)

    return file_path
