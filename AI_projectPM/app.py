import streamlit as st
import requests
import json
import os
from io import BytesIO
from docx import Document

# =========================
# 页面配置
# =========================
st.set_page_config(
    page_title="AI PM Agent",
    page_icon="🚀",
    layout="wide"
)

# =========================
# UI 样式（重点升级）
# =========================
st.markdown("""
<style>
.main-title {
    font-size: 42px;
    font-weight: 800;
    margin-bottom: 5px;
}
.sub-title {
    color: #666;
    margin-bottom: 30px;
}

.block {
    background: #FFFFFF;
    padding: 20px;
    border-radius: 14px;
    border: 1px solid #eee;
    margin-bottom: 15px;
}

.metric-card {
    background: #000;
    color: white;
    padding: 16px;
    border-radius: 12px;
    text-align: center;
}

.card-title {
    font-weight: 700;
    font-size: 18px;
    margin-bottom: 10px;
}

.result-card {
    background: #F7F7F7;
    padding: 18px;
    border-radius: 12px;
    line-height: 1.7;
    white-space: pre-wrap;
}

button[kind="primary"] {
    background-color: black !important;
    border-radius: 10px !important;
}
</style>
""", unsafe_allow_html=True)

# =========================
# 顶部
# =========================
st.markdown('<div class="main-title">AI Product Manager Agent</div>', unsafe_allow_html=True)
st.markdown('<div class="sub-title">一键生成 PRD / 竞品分析 / 产品机会点</div>', unsafe_allow_html=True)

# =========================
# Sidebar（产品感增强）
# =========================
with st.sidebar:
    st.markdown("## ⚙️ 系统配置")

    workflow_url = st.text_input(
        "Coze 工作流地址",
        value="https://7fv2jsrt7q.coze.site/run"
    )

    st.markdown("---")
    st.markdown("## 📊 使用说明")
    st.write("1. 输入产品名称")
    st.write("2. 点击生成")
    st.write("3. 下载报告")

    st.markdown("---")
    st.markdown("## 🔐 Secrets")
    st.code('COZE_API_TOKEN = "xxx"')

# =========================
# 工具函数
# =========================
def get_token():
    return st.secrets.get("COZE_API_TOKEN", os.getenv("COZE_API_TOKEN", ""))

def generate_word(text):
    doc = Document()
    doc.add_heading("AI 产品分析报告", 0)
    doc.add_paragraph(text)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def call_api(product):
    token = get_token()
    headers = {
        "Authorization": f"Bearer {token}",
        "Content-Type": "application/json"
    }

    payload = {"query": product}

    with st.spinner("AI 正在分析中..."):
        res = requests.post(workflow_url, headers=headers, json=payload, timeout=120)

    if res.status_code == 200:
        try:
            return res.json()
        except:
            return {"output": res.text}
    else:
        st.error("请求失败")
        return None

# =========================
# 输入区（产品化）
# =========================
col1, col2 = st.columns([3,1])

with col1:
    product = st.text_input("📦 输入产品名称", placeholder="例如：Keep / 小红书 / ChatGPT")

with col2:
    run = st.button("生成分析", use_container_width=True)

# =========================
# 数据统计（增强产品感）
# =========================
st.markdown("### 📊 系统状态")

col1, col2, col3 = st.columns(3)
col1.markdown('<div class="metric-card">⚡ AI引擎<br><b>Online</b></div>', unsafe_allow_html=True)
col2.markdown('<div class="metric-card">📈 分析能力<br><b>PRD + 竞品</b></div>', unsafe_allow_html=True)
col3.markdown('<div class="metric-card">⏱ 响应时间<br><b>30-90s</b></div>', unsafe_allow_html=True)

# =========================
# 主逻辑
# =========================
if run:
    if not product:
        st.warning("请输入产品名称")
    else:
        result = call_api(product)

        if result:
            st.success("分析完成")

            data = result.get("data") or result.get("output") or result

            if isinstance(data, str):
                try:
                    data = json.loads(data)
                except:
                    pass

            positioning = data.get("positioning_analysis", "")
            comp = data.get("competitive_analysis", "")
            prd = data.get("prd_document", "")
            opp = data.get("opportunity_analysis", "")

            st.markdown("## 📊 分析结果")

            col1, col2 = st.columns(2)

            with col1:
                st.markdown('<div class="block"><div class="card-title">🎯 产品定位</div>', unsafe_allow_html=True)
                st.markdown(f'<div class="result-card">{positioning}</div></div>', unsafe_allow_html=True)

                st.markdown('<div class="block"><div class="card-title">⚔️ 竞品分析</div>', unsafe_allow_html=True)
                st.markdown(f'<div class="result-card">{comp}</div></div>', unsafe_allow_html=True)

            with col2:
                st.markdown('<div class="block"><div class="card-title">📄 PRD 文档</div>', unsafe_allow_html=True)
                st.markdown(f'<div class="result-card">{prd}</div></div>', unsafe_allow_html=True)

                st.markdown('<div class="block"><div class="card-title">💡 机会分析</div>', unsafe_allow_html=True)
                st.markdown(f'<div class="result-card">{opp}</div></div>', unsafe_allow_html=True)

            # 下载
            st.markdown("### 📥 下载报告")
            col1, col2 = st.columns(2)

            col1.download_button(
                "下载 Word",
                data=generate_word(prd),
                file_name=f"{product}_PRD.docx"
            )

            col2.download_button(
                "下载 JSON",
                data=json.dumps(result, ensure_ascii=False),
                file_name=f"{product}.json"
            )
