from docx import Document

def generate_docx(doc_data, product_name):
    document = Document()

    # 标题
    document.add_heading(f"{product_name} PRD", 0)

    for section in doc_data:
        document.add_heading(section["title"], level=1)
        document.add_paragraph(section["content"])

    file_path = f"/tmp/{product_name}_PRD.docx"
    document.save(file_path)

    return file_path
